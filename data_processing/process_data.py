# -*- coding: utf-8 -*-
import os
import json
import time
from pathlib import Path

import openai
from dotenv import load_dotenv
from docx import Document

DATA_PROCESSING_DIR = Path(__file__).parent
load_dotenv(DATA_PROCESSING_DIR.parent / '.env')  # load environment variables from .env file

# Set up OpenAI API credentials
openai.api_key = os.getenv('OPENAI_API_KEY')


# openai.api_key = "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"


# Define function for processing a single post
def process_post(post_path):
    # Load JSON data from file
    with open(post_path) as f:
        post_data = json.load(f)

    # Extract title and content from JSON data
    content = post_data[0]['content']

    # Call OpenAI GPT-3 to generate a title
    title_prompt = f"Generate a title based on the content:\n{content}"
    title_response = openai.Completion.create(
        engine="text-davinci-002",
        prompt=title_prompt,
        max_tokens=60,
        n=1,
        stop=None,
        temperature=0.7,
    )
    title = title_response.choices[0].text.strip()

    # Call OpenAI GPT-3 to summarize the content within 800 characters
    summary_prompt = f"create the post sized from **750 to 850** characters based on the text below. Shorten it while retaining **key metrics and numbers**, divide it to paragraphs and leave 1 blank line between paragraphs:\n{content}"
    summary_response = openai.Completion.create(
        engine="text-davinci-002",
        prompt=summary_prompt,
        max_tokens=1300,
        n=1,
        stop=None,
        temperature=0.7,
    )
    summary = summary_response.choices[0].text.strip()
    # Format the text for Telegram
    doc = Document()
    title = f"**{title}**"
    doc.add_heading(title, level=1)
    doc.add_paragraph('_' * 30)  # Add line of underscores
    doc.add_paragraph()  # Add blank line after title
    paragraphs = summary.split("\n\n")  # Split text into paragraphs

    for i, paragraph in enumerate(paragraphs):
        if paragraph.strip():  # Check if paragraph is non-empty
            if i > 0:
                doc.add_paragraph()  # Add completely blank line between paragraphs
            p = doc.add_paragraph("• ")  # Start paragraph with symbol
            p.add_run(paragraph.strip())
        else:
            doc.add_paragraph()  # Add completely blank line between paragraphs

    # Save the formatted text as a docx file in the "Ready_posts" folder
    processed_filename = DATA_PROCESSING_DIR / f"Ready_posts/{os.path.splitext(os.path.basename(post_path))[0]}.docx"

    doc.save(processed_filename)

    # Move original JSON file to "Processed" folder
    processed_path = DATA_PROCESSING_DIR / f"Processed/{os.path.basename(post_path)}"
    os.rename(post_path, processed_path)


# Run indefinitely
while True:
    # Check for new JSON files in the "Posts" folder
    posts_path = DATA_PROCESSING_DIR / 'Posts'

    for filename in os.listdir(posts_path):
        if filename.endswith(".json"):
            post_path = posts_path / filename

            process_post(post_path)

    # Wait for 3 minutes before checking again
    time.sleep(3 * 60)
